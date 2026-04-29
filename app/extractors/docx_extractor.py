from __future__ import annotations

from pathlib import Path

from app.core.exceptions import ExtractionError
from app.extractors.base_extractor import BaseExtractor


class DocxExtractor(BaseExtractor):
    def extract(self) -> str:
        try:
            from docx import Document
            doc = Document(str(self.file_path))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            return "\n".join(paragraphs)
        except Exception as exc:
            raise ExtractionError(f"Failed to extract DOCX: {exc}") from exc
